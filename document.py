from docx import Document

d = Document()
d.add_heading('Не забыть взять с собой', 0)
d.add_paragraph('Документы\n\nКлючи')

p = d.add_paragraph('Ноутбук ', style='List Bullet')
p.add_run('с зарядкой').bold = True

d.add_paragraph('Косметика\n', style='List Number')
d.add_paragraph('Плойка', style='List Number')

d.add_paragraph('Конец списка', style='Intense Quote')

table = d.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = '1'
table.rows[0].cells[1].text = '2'

d.save('text.docx')
